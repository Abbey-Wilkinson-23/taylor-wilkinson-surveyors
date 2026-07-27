"""
Parse the TWS postcode coverage Word document and generate a standalone HTML tool.

Fee classification:
  🟢 Green (STANDARD) - Surveyor text is green (preferred, standard rates), OR no special
                         colour/text and table header has no fee warning.
  🟡 Amber (HIGHER)   - Table header has FIX FEE / MIN FEE / HIGHER FEE and surveyor is
                         not green-coloured, OR blue-coloured surveyor, OR fee keyword in text.
  🔴 Red (QUOTABLE)   - Coverage text contains "Quotable work only" (and surveyor is not blue).
"""

import re
import json
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = Path("/Users/abbeywilkinson./Dropbox/MARKETING TAYLOR WILKINSON/Postcode Lists/Members postcode coverage internal July 2026.docx")
OUT_PATH = Path("/Users/abbeywilkinson./dev/tws/postcode_coverage.html")

# Blue colour = explicitly HIGHER FEE (document note: "BLUE – HIGHER FEES")
BLUE_COLOURS  = {"0070C0", "00B0F0", "1F4E79", "2E75B6", "4472C4"}
# Green colour = preferred surveyor at standard rates ("GREEN # - PREFERRED SURVEYOR TO USE")
GREEN_COLOURS = {"00B050", "70AD47", "92D050", "A8D08D"}

# Keywords that signal QUOTABLE or HIGHER in coverage/name text
QUOTABLE_RE = re.compile(r"quotable work only|quotable only", re.I)
HIGHER_RE   = re.compile(r"higher fee|higher fees|slightly higher fees|fix fee|min fee|cancellation charges", re.I)


def get_run_colours(cell):
    """Return set of hex colour values from all runs in a cell (excluding 'auto' / None)."""
    colours = set()
    for para in cell.paragraphs:
        for run in para.runs:
            rPr = run._r.find(qn("w:rPr"))
            if rPr is not None:
                col_el = rPr.find(qn("w:color"))
                if col_el is not None:
                    val = col_el.get(qn("w:val"), "")
                    if val and val.upper() not in ("AUTO", "000000", "FFFFFF", ""):
                        colours.add(val.upper())
    return colours


def classify(table_higher: bool, name: str, coverage: str, name_colours: set) -> str:
    """Return 'HIGHER', 'QUOTABLE', or 'STANDARD'."""
    has_quotable     = bool(QUOTABLE_RE.search(coverage))
    has_higher_text  = bool(HIGHER_RE.search(coverage)) or bool(HIGHER_RE.search(name))
    is_blue          = bool(name_colours & BLUE_COLOURS)
    is_green         = bool(name_colours & GREEN_COLOURS)

    # Blue text = explicitly higher fees (overrides everything)
    if is_blue or has_higher_text:
        return "HIGHER"
    # Green text = preferred surveyor, works to standard scale
    # (table-level FIX/MIN FEE does not apply to them)
    if is_green:
        return "QUOTABLE" if has_quotable else "STANDARD"
    # Black/uncoloured text: fall back to table-level and coverage text
    if has_quotable:
        return "QUOTABLE"
    if table_higher:
        return "HIGHER"
    return "STANDARD"


def parse_doc():
    doc = Document(DOCX_PATH)
    entries = []
    last_valid_area = ""

    for table in doc.tables:
        if not table.rows:
            continue

        # Header row: first cell text is the postcode area code (possibly with fee suffix)
        header_text = table.rows[0].cells[0].text.strip()
        first_word  = header_text.split()[0].upper() if header_text else ""

        # Tables whose header starts with # have a surveyor name instead of a postcode area.
        # Inherit the postcode area from the previous valid table and treat ALL rows as surveyors.
        if first_word.startswith("#"):
            postcode_area = last_valid_area
            table_higher  = False
            surveyor_rows = table.rows          # first row is a surveyor, not a header
        else:
            postcode_area = first_word
            if not postcode_area or len(postcode_area) > 4:
                continue  # skip unrecognised tables
            last_valid_area = postcode_area
            table_higher    = bool(HIGHER_RE.search(header_text))
            surveyor_rows   = table.rows[1:]    # first row is the postcode header

        for row in surveyor_rows:
            cells = row.cells
            if len(cells) < 2:
                continue

            name_cell     = cells[0]
            coverage_cell = cells[1]
            work_cell     = cells[2] if len(cells) > 2 else None

            name_text     = name_cell.text.strip()
            coverage_text = coverage_cell.text.strip()
            work_text     = work_cell.text.strip() if work_cell else ""

            if not name_text:
                continue  # skip empty rows

            # Detect preferred (#/##)
            preferred = "##" if name_text.startswith("##") else ("#" if name_text.startswith("#") else "")
            clean_name = name_text.lstrip("#").strip()

            name_colours = get_run_colours(name_cell)
            fee_cat = classify(table_higher, name_text, coverage_text, name_colours)

            entries.append({
                "postcode_area": postcode_area,
                "name":          clean_name,
                "preferred":     preferred,
                "coverage":      coverage_text,
                "work_types":    work_text,
                "fee_cat":       fee_cat,
                "table_higher":  table_higher,
            })

    return entries


# ---------------------------------------------------------------------------
# HTML generation
# ---------------------------------------------------------------------------

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>TWS Postcode Coverage</title>
<style>
  * { box-sizing: border-box; margin: 0; padding: 0; }
  html, body { height: 100%; }
  body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; background: #f5f5f5; color: #222; display: flex; flex-direction: column; height: 100vh; overflow: hidden; }
  .table-wrapper { flex: 1; overflow-y: auto; }
  header { background: #794899; color: #fff; padding: 12px 24px; display: flex; align-items: center; gap: 16px; }
  header h1 { font-size: 1.25rem; font-weight: 600; }
  header small { opacity: .75; font-size: .8rem; }

  .controls { background: #fff; border-bottom: 1px solid #ddd; padding: 10px 24px; display: flex; gap: 10px; flex-wrap: wrap; align-items: center; }
  .controls input, .controls select { padding: 7px 10px; border: 1px solid #ccc; border-radius: 6px; font-size: .9rem; outline: none; }
  .controls input:focus, .controls select:focus { border-color: #794899; }
  .controls input { min-width: 220px; flex: 1; }
  .legend { margin-left: auto; display: flex; gap: 12px; align-items: center; font-size: .82rem; }
  .dot { display: inline-block; width: 10px; height: 10px; border-radius: 50%; margin-right: 4px; }
  .dot.g { background: #389e0d; }
  .dot.a { background: #d48806; }
  .dot.r { background: #cf1322; }

  .count-bar { background: #fff; padding: 5px 24px; border-bottom: 1px solid #eee; font-size: .82rem; color: #666; }

  table { width: 100%; border-collapse: collapse; background: #fff; }
  thead th { background: #f0f0f0; padding: 8px 12px; text-align: left; font-size: .8rem; color: #555; border-bottom: 2px solid #ddd; position: sticky; top: 0; z-index: 10; cursor: pointer; user-select: none; white-space: nowrap; }
  thead th:hover { background: #e8e0f0; }
  thead th.no-sort { cursor: default; }
  thead th.no-sort:hover { background: #f0f0f0; }
  thead th .sort-arrow { margin-left: 4px; opacity: .4; }
  thead th.asc .sort-arrow::after  { content: " ▲"; opacity: 1; }
  thead th.desc .sort-arrow::after { content: " ▼"; opacity: 1; }
  tbody tr { border-bottom: 1px solid #f0f0f0; transition: background .1s; }
  tbody tr:hover { background: #fafafa; }
  tbody tr.custom-row { background: #f9f0ff; }
  tbody tr.custom-row:hover { background: #f0e0ff; }
  td { padding: 7px 12px; font-size: .85rem; vertical-align: top; }
  td.postcode { font-weight: 700; font-size: .9rem; color: #794899; min-width: 60px; }
  td.name { min-width: 200px; }
  td.coverage { color: #444; min-width: 180px; max-width: 320px; white-space: pre-line; }
  td.work { color: #555; min-width: 90px; }
  td.actions { width: 40px; text-align: center; }
  .badge { display: inline-flex; align-items: center; gap: 5px; padding: 2px 8px; border-radius: 12px; font-size: .76rem; font-weight: 600; }
  .badge.STANDARD { background: #f6ffed; color: #389e0d; border: 1px solid #b7eb8f; }
  .badge.QUOTABLE { background: #fff1f0; color: #cf1322; border: 1px solid #ffa39e; }
  .badge.HIGHER   { background: #fffbe6; color: #d48806; border: 1px solid #ffe58f; }
  .pref { font-size: .7rem; color: #794899; font-weight: 700; margin-left: 4px; }
  .custom-tag { font-size: .68rem; background: #794899; color: #fff; border-radius: 4px; padding: 1px 5px; margin-left: 5px; }

  .btn-add { background: #794899; color: #fff; border: none; border-radius: 6px; padding: 7px 14px; font-size: .9rem; cursor: pointer; white-space: nowrap; }
  .btn-add:hover { background: #5e3478; }
  .btn-secondary { background: #f5f5f5; color: #444; border: 1px solid #ccc; border-radius: 6px; padding: 7px 14px; font-size: .9rem; cursor: pointer; white-space: nowrap; }
  .btn-secondary:hover { background: #eee; border-color: #aaa; }
  .btn-del { background: none; border: none; cursor: pointer; color: #bbb; font-size: 1rem; padding: 2px 6px; border-radius: 4px; line-height: 1; }
  .btn-del:hover { color: #cf1322; background: #fff1f0; }

  /* Modal */
  .modal-backdrop { display: none; position: fixed; inset: 0; background: rgba(0,0,0,.4); z-index: 100; align-items: center; justify-content: center; }
  .modal-backdrop.open { display: flex; }
  .modal { background: #fff; border-radius: 10px; padding: 28px; width: 480px; max-width: 95vw; box-shadow: 0 8px 32px rgba(0,0,0,.18); }
  .modal h2 { font-size: 1.1rem; margin-bottom: 20px; color: #794899; }
  .form-row { margin-bottom: 14px; }
  .form-row label { display: block; font-size: .82rem; color: #555; margin-bottom: 4px; font-weight: 500; }
  .form-row input, .form-row select, .form-row textarea { width: 100%; padding: 8px 10px; border: 1px solid #ccc; border-radius: 6px; font-size: .9rem; font-family: inherit; outline: none; }
  .form-row input:focus, .form-row select:focus, .form-row textarea:focus { border-color: #794899; }
  .form-row textarea { resize: vertical; min-height: 60px; }
  .modal-actions { display: flex; justify-content: flex-end; gap: 10px; margin-top: 20px; }
  .btn-cancel { background: #f5f5f5; border: 1px solid #ddd; border-radius: 6px; padding: 8px 18px; cursor: pointer; font-size: .9rem; }
  .btn-cancel:hover { background: #eee; }
  .btn-save { background: #794899; color: #fff; border: none; border-radius: 6px; padding: 8px 18px; cursor: pointer; font-size: .9rem; }
  .btn-save:hover { background: #5e3478; }

  .no-results { text-align: center; padding: 40px; color: #aaa; font-size: 1rem; }
  @media (max-width: 700px) {
    .legend { display: none; }
    .controls { gap: 6px; }
    td, th { padding: 5px 8px; }
  }
</style>
</head>
<body>
<header>
  <div>
    <h1>TWS Postcode Coverage</h1>
    <small>Internal use only — Members coverage July 2026</small>
  </div>
  <div class="legend">
    <span><span class="dot g"></span>Standard fee scale</span>
    <span><span class="dot r"></span>Quotable work only</span>
    <span><span class="dot a"></span>Higher / fix / min fee</span>
  </div>
</header>

<div class="controls">
  <input type="text" id="search" placeholder="Search by surveyor name, firm, coverage…" oninput="applyFilters()">
  <select id="postcodeFilter" onchange="applyFilters()">
    <option value="">All postcode areas</option>
    __POSTCODE_OPTIONS__
  </select>
  <select id="feeFilter" onchange="applyFilters()">
    <option value="">All fee types</option>
    <option value="STANDARD">Standard fee scale</option>
    <option value="QUOTABLE">Quotable work only</option>
    <option value="HIGHER">Higher / fix / min fee</option>
  </select>
  <select id="workFilter" onchange="applyFilters()">
    <option value="">All work types</option>
    <option value="R">Residential (R)</option>
    <option value="C">Commercial (C)</option>
    <option value="GDV">GDV</option>
  </select>
  <button class="btn-add" onclick="openModal()">+ Add Surveyor</button>
  <button class="btn-secondary" onclick="exportChanges()" title="Export your additions and removals to share with others">Export Changes</button>
  <label class="btn-secondary" title="Import changes exported by someone else">
    Import Changes
    <input type="file" id="importFile" accept=".json" style="display:none" onchange="importChanges(event)">
  </label>
</div>
<div class="count-bar" id="countBar"></div>

<div class="table-wrapper">
<table>
  <thead>
    <tr>
      <th onclick="sortBy('postcode_area')" data-col="postcode_area">Area <span class="sort-arrow"></span></th>
      <th onclick="sortBy('name')" data-col="name">Surveyor / Firm <span class="sort-arrow"></span></th>
      <th onclick="sortBy('coverage')" data-col="coverage">Coverage <span class="sort-arrow"></span></th>
      <th onclick="sortBy('work_types')" data-col="work_types">Work Types <span class="sort-arrow"></span></th>
      <th onclick="sortBy('fee_cat')" data-col="fee_cat">Fee Category <span class="sort-arrow"></span></th>
      <th class="no-sort"></th>
    </tr>
  </thead>
  <tbody id="tbody"></tbody>
</table>
<div class="no-results" id="noResults" style="display:none">No matching surveyors found.</div>
</div>

<!-- Add Surveyor Modal -->
<div class="modal-backdrop" id="modal" onclick="e => { if(e.target===this) closeModal(); }">
  <div class="modal" onclick="event.stopPropagation()">
    <h2>Add Surveyor</h2>
    <div class="form-row">
      <label>Postcode Area *</label>
      <input id="f-area" type="text" placeholder="e.g. SW" style="text-transform:uppercase">
    </div>
    <div class="form-row">
      <label>Surveyor / Firm Name *</label>
      <input id="f-name" type="text" placeholder="e.g. John Smith (123) Acme Surveyors">
    </div>
    <div class="form-row">
      <label>Coverage</label>
      <textarea id="f-coverage" placeholder="e.g. 1-10, 15, 20-25 (check distance)"></textarea>
    </div>
    <div class="form-row">
      <label>Work Types</label>
      <input id="f-work" type="text" placeholder="e.g. R / C / GDV">
    </div>
    <div class="form-row">
      <label>Fee Category *</label>
      <select id="f-fee">
        <option value="STANDARD">Standard fee scale</option>
        <option value="QUOTABLE">Quotable work only</option>
        <option value="HIGHER">Higher / fix / min fee</option>
      </select>
    </div>
    <div class="modal-actions">
      <button class="btn-cancel" onclick="closeModal()">Cancel</button>
      <button class="btn-save" onclick="saveSurveyor()">Add Surveyor</button>
    </div>
  </div>
</div>

<script>
const BASE_DATA = __DATA__;

const FEE_LABELS = {
  STANDARD: "Standard",
  QUOTABLE: "Quotable only",
  HIGHER:   "Higher fee",
};

const LS_CUSTOM  = "tws_custom_surveyors";
const LS_REMOVED = "tws_removed_surveyors";

function loadCustom()  { try { return JSON.parse(localStorage.getItem(LS_CUSTOM)  || "[]"); } catch { return []; } }
function loadRemoved() { try { return JSON.parse(localStorage.getItem(LS_REMOVED) || "[]"); } catch { return []; } }
function saveCustom(d)  { localStorage.setItem(LS_CUSTOM,  JSON.stringify(d)); }
function saveRemoved(d) { localStorage.setItem(LS_REMOVED, JSON.stringify(d)); }

function buildData() {
  const removed = new Set(loadRemoved());
  const base = BASE_DATA
    .map((r, i) => ({ ...r, _id: "b" + i, _custom: false }))
    .filter(r => !removed.has(r._id));
  const custom = loadCustom().map(r => ({ ...r, _custom: true }));
  return [...base, ...custom];
}

let sortCol = "postcode_area";
let sortDir = 1;

function sortBy(col) {
  if (sortCol === col) { sortDir *= -1; }
  else { sortCol = col; sortDir = 1; }
  document.querySelectorAll("thead th[data-col]").forEach(th => {
    th.classList.remove("asc", "desc");
    if (th.dataset.col === col) th.classList.add(sortDir === 1 ? "asc" : "desc");
  });
  applyFilters();
}

function applyFilters() {
  const q        = document.getElementById("search").value.trim().toUpperCase();
  const postcode = document.getElementById("postcodeFilter").value;
  const fee      = document.getElementById("feeFilter").value;
  const work     = document.getElementById("workFilter").value;

  let rows = buildData().filter(r => {
    if (postcode && r.postcode_area !== postcode) return false;
    if (fee && r.fee_cat !== fee) return false;
    if (work && !r.work_types.toUpperCase().includes(work)) return false;
    if (!q) return true;
    return (
      r.postcode_area.includes(q) ||
      r.name.toUpperCase().includes(q) ||
      r.coverage.toUpperCase().includes(q)
    );
  });

  rows.sort((a, b) => {
    const av = a[sortCol] || "";
    const bv = b[sortCol] || "";
    return av < bv ? -sortDir : av > bv ? sortDir : 0;
  });

  const tbody = document.getElementById("tbody");
  tbody.innerHTML = rows.map(r => {
    const prefBadge   = r.preferred ? `<span class="pref">${esc(r.preferred)}</span>` : "";
    const customTag   = r._custom ? `<span class="custom-tag">added</span>` : "";
    return `<tr class="${r._custom ? "custom-row" : ""}">
      <td class="postcode">${esc(r.postcode_area)}</td>
      <td class="name">${esc(r.name)}${prefBadge}${customTag}</td>
      <td class="coverage">${esc(r.coverage)}</td>
      <td class="work">${esc(r.work_types)}</td>
      <td><span class="badge ${r.fee_cat}">${FEE_LABELS[r.fee_cat]}</span></td>
      <td class="actions"><button class="btn-del" title="Remove" onclick="removeSurveyor('${esc(r._id)}')">&#x2715;</button></td>
    </tr>`;
  }).join("");

  document.getElementById("countBar").textContent =
    `${rows.length.toLocaleString()} surveyor${rows.length !== 1 ? "s" : ""} shown`;
  document.getElementById("noResults").style.display = rows.length ? "none" : "block";
}

function removeSurveyor(id) {
  if (!confirm("Remove this surveyor from the list?")) return;
  if (id.startsWith("b")) {
    const removed = loadRemoved();
    removed.push(id);
    saveRemoved(removed);
  } else {
    const custom = loadCustom().filter(r => r._id !== id);
    saveCustom(custom);
  }
  applyFilters();
}

function openModal() {
  document.getElementById("f-area").value     = "";
  document.getElementById("f-name").value     = "";
  document.getElementById("f-coverage").value = "";
  document.getElementById("f-work").value     = "";
  document.getElementById("f-fee").value      = "STANDARD";
  document.getElementById("modal").classList.add("open");
  document.getElementById("f-area").focus();
}

function closeModal() {
  document.getElementById("modal").classList.remove("open");
}

function saveSurveyor() {
  const area = document.getElementById("f-area").value.trim().toUpperCase();
  const name = document.getElementById("f-name").value.trim();
  if (!area || !name) { alert("Postcode area and name are required."); return; }

  const entry = {
    _id:          "c" + Date.now(),
    _custom:      true,
    postcode_area: area,
    name:          name,
    preferred:     "",
    coverage:      document.getElementById("f-coverage").value.trim(),
    work_types:    document.getElementById("f-work").value.trim(),
    fee_cat:       document.getElementById("f-fee").value,
    table_higher:  false,
  };

  const custom = loadCustom();
  custom.push(entry);
  saveCustom(custom);
  closeModal();
  applyFilters();
}

// Export changes as a JSON file
function exportChanges() {
  const custom  = loadCustom();
  const removed = loadRemoved();
  if (!custom.length && !removed.length) {
    alert("No changes to export yet. Add or remove surveyors first.");
    return;
  }
  const payload = { version: 1, custom, removed, exported_at: new Date().toISOString() };
  const blob = new Blob([JSON.stringify(payload, null, 2)], { type: "application/json" });
  const a = document.createElement("a");
  a.href = URL.createObjectURL(blob);
  a.download = "tws_postcode_changes.json";
  a.click();
  URL.revokeObjectURL(a.href);
}

// Import changes from a JSON file (merges with existing local changes)
function importChanges(event) {
  const file = event.target.files[0];
  if (!file) return;
  const reader = new FileReader();
  reader.onload = function(e) {
    try {
      const data = JSON.parse(e.target.result);
      if (data.version !== 1 || !Array.isArray(data.custom) || !Array.isArray(data.removed)) {
        alert("Invalid changes file.");
        return;
      }
      // Merge: add any custom entries not already present (by _id)
      const existingCustom = loadCustom();
      const existingIds    = new Set(existingCustom.map(r => r._id));
      const newEntries     = data.custom.filter(r => !existingIds.has(r._id));
      saveCustom([...existingCustom, ...newEntries]);
      // Merge removed lists
      const existingRemoved = new Set(loadRemoved());
      data.removed.forEach(id => existingRemoved.add(id));
      saveRemoved([...existingRemoved]);
      applyFilters();
      const msg = [];
      if (newEntries.length) msg.push(`${newEntries.length} new surveyor(s) added`);
      if (data.removed.length) msg.push(`${data.removed.length} removal(s) applied`);
      alert("Changes imported: " + (msg.join(", ") || "nothing new to merge") + ".");
    } catch {
      alert("Could not read the file. Make sure it's a valid changes export.");
    }
    event.target.value = "";
  };
  reader.readAsText(file);
}

// Close modal on backdrop click
document.getElementById("modal").addEventListener("click", function(e) {
  if (e.target === this) closeModal();
});

// Close modal on Escape
document.addEventListener("keydown", e => { if (e.key === "Escape") closeModal(); });

// Initial render
sortBy("postcode_area");
</script>
</body>
</html>
"""


def main():
    print("Parsing document…")
    entries = parse_doc()
    print(f"Parsed {len(entries)} entries across {len(set(e['postcode_area'] for e in entries))} postcode areas")

    by_cat = {}
    for e in entries:
        by_cat[e["fee_cat"]] = by_cat.get(e["fee_cat"], 0) + 1
    print("Fee categories:", by_cat)

    areas = sorted(set(e["postcode_area"] for e in entries))
    postcode_options = "\n    ".join(
        f'<option value="{a}">{a}</option>' for a in areas
    )

    data_json = json.dumps(entries, ensure_ascii=False)
    html = HTML_TEMPLATE.replace("__DATA__", data_json).replace("__POSTCODE_OPTIONS__", postcode_options)

    OUT_PATH.write_text(html, encoding="utf-8")
    print(f"Output written to {OUT_PATH}")


if __name__ == "__main__":
    main()
